# Python import
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# Local import
from core.utils import format_paragraph


class ReportDocument:
    def __init__(self, report):
        self.report = report
        self.document = self.__generate(report)

    def save(self, stream):
        self.document.save(stream)
        stream.seek(0)

    def name(self):
        return self.report.title

    @staticmethod
    def __get_value(value):
        return '' if value is None else value

    @staticmethod
    def __generate(report):
        document = Document()

        report_heading = document.add_heading(f'{report.title}', 0)
        report_heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        format_paragraph(document.add_paragraph(), f'<b>Open Date</b> : {report.open_date}')
        format_paragraph(document.add_paragraph(), f'<b>Close Date</b> : {report.close_date}')
        format_paragraph(document.add_paragraph(), f'<b>Type</b> :  {report.type}')
        format_paragraph(document.add_paragraph(), f'<b>Descriptions</b> : \n {ReportDocument.__get_value(report.description)}')

        for player_report in report.player_reports.all():
            ReportDocument.__fill_player_report(document, player_report)

        return document

    @staticmethod
    def __fill_player_report(document, player_report):
        player = player_report.player

        player_report_heading = document.add_heading(f'{player.name} {player.surname}', 1)
        player_report_heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        format_paragraph(document.add_paragraph(), f'<b>Name</b> : {player.surname}')
        format_paragraph(document.add_paragraph(), f'<b>Surname</b> : {player.surname}')
        format_paragraph(document.add_paragraph(), f'<b>Position</b> : {player.position}')
        format_paragraph(document.add_paragraph(), f'<b>Birth Date</b> : {player.birth_date}')
        format_paragraph(document.add_paragraph(), f'<b>Country</b> : {player.country}')
        format_paragraph(document.add_paragraph(), f'<b>Dominant Leg</b> : {player.dominant_leg}')
        format_paragraph(document.add_paragraph(), f'<b>team</b> : {player.team.name}')
        format_paragraph(document.add_paragraph(), f'<b>league</b> : {player.league.name}')
        format_paragraph(document.add_paragraph(), f'<b>Descriptions</b> : \n {ReportDocument.__get_value(player_report.description)}')

        for profile in player_report.profiles.all():
            ReportDocument.__fill_profile(document, profile)

    @staticmethod
    def __fill_statistic(statistics_table, statistics):
        for statistic in statistics:
            statistic_cells = statistics_table.add_row().cells
            format_paragraph(statistic_cells[0].paragraphs[0], f'{statistic.name}', font_size=8)
            format_paragraph(statistic_cells[1].paragraphs[0], f'{ReportDocument.__get_value(statistic.value)}', font_size=8)

    @staticmethod
    def __fill_profile(document, profile):
        document.add_heading(f'{profile.name} :', 2)
        format_paragraph(document.add_paragraph(), f'<b>Descriptions</b> : \n {ReportDocument.__get_value(profile.description)}')

        groups = profile.groups.all()
        statistics = profile.statistics.all()

        statistics_table = document.add_table(rows=1, cols=2)
        statistics_table.style = 'Table Grid'

        heading_cells = statistics_table.rows[0].cells
        format_paragraph(heading_cells[0].paragraphs[0], '<b>Name</b>', font_size=12)
        format_paragraph(heading_cells[1].paragraphs[0], '<b>Value</b>', font_size=12)

        for group in groups:
            group_cells = statistics_table.add_row().cells
            group_cells[0].merge(group_cells[1])
            group_cells[0].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            format_paragraph(group_cells[0].paragraphs[0], f'<i>{group.name}</i> :', font_size=10)
            ReportDocument.__fill_statistic(statistics_table, group.statistics.all())

        ReportDocument.__fill_statistic(statistics_table, statistics)
