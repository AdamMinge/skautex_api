# Python import
import os
import re
from docx import Document
# Local import
from contact_details.models import ContactDetailTypeChoice


def create_invitation_file_path(instance, filename):
    return os.path.join(
        'players/invitations',
        filename
    )


class InvitationDocument:
    def __init__(self, invitation_template, player, owner, trainer, creation_date, start_date, end_date):
        self.invitation_template = invitation_template
        self.player = player
        self.owner = owner
        self.trainer = trainer
        self.creation_date = creation_date
        self.start_date = start_date
        self.end_date = end_date
        self.document = self.__generate()

    def save(self, stream):
        self.document.save(stream)
        stream.seek(0)

    def name(self):
        return f'invitation {self.player.name} {self.player.surname}'

    @property
    def invitation_creation_date(self):
        return f'{self.creation_date.date()}'

    @property
    def invitation_player(self):
        return f'{self.player.name} {self.player.surname} ur. {self.player.birth_date.year} r.'

    @property
    def invitation_date(self):
        return f'{self.start_date} - {self.end_date}'

    @property
    def invitation_trainer(self):
        contact_detail = self.trainer.contact_details.filter(type=ContactDetailTypeChoice.PHONE).first()
        phone_number = contact_detail.value if contact_detail is not None else '<missing_phone_number>'
        name = self.trainer.first_name if self.trainer.first_name else '<missing_name>'
        surname = self.trainer.last_name if self.trainer.last_name else '<missing_surname>'

        return f'{name} {surname} {phone_number}'

    @property
    def invitation_owner(self):
        name = self.owner.first_name if self.owner.first_name else '<missing_name>'
        surname = self.owner.last_name if self.owner.last_name else '<missing_surname>'
        return f'{name} {surname}'

    def __generate(self):
        document = Document(self.invitation_template.template)

        InvitationDocument.docx_replace_regex(
            document, re.compile(r"<invitation_creation_date>"), self.invitation_creation_date)
        InvitationDocument.docx_replace_regex(
            document, re.compile(r"<invitation_player>"), self.invitation_player)
        InvitationDocument.docx_replace_regex(
            document, re.compile(r"<invitation_date>"), self.invitation_date)
        InvitationDocument.docx_replace_regex(
            document, re.compile(r"<invitation_trainer>"), self.invitation_trainer)
        InvitationDocument.docx_replace_regex(
            document, re.compile(r"<invitation_owner>"), self.invitation_owner)

        return document

    @staticmethod
    def docx_replace_regex(doc_obj, regex, replace):
        for p in doc_obj.paragraphs:
            print(p.text)
            if regex.search(p.text):
                inline = p.runs
                for i in range(len(inline)):
                    if regex.search(inline[i].text):
                        text = regex.sub(replace, inline[i].text)
                        inline[i].text = text

        for table in doc_obj.tables:
            for row in table.rows:
                for cell in row.cells:
                    InvitationDocument.docx_replace_regex(cell, regex, replace)

